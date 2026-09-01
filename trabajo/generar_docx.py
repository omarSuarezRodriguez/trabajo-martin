# -*- coding: utf-8 -*-
"""Genera ANALISIS_ARQUITECTURA.docx a partir del markdown y cuenta palabras."""

from __future__ import annotations

import re
from pathlib import Path

from docx import Document
from docx.enum.section import WD_SECTION
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.enum.style import WD_STYLE_TYPE
from docx.oxml import OxmlElement
from docx.oxml.ns import qn, nsmap
from docx.shared import Cm, Pt, RGBColor, Inches, Emu, Twips
from docx.opc.constants import RELATIONSHIP_TYPE as RT

ROOT = Path(__file__).resolve().parent
MD_PATH = ROOT / "ANALISIS_ARQUITECTURA.md"
DOCX_PATH = ROOT / "ANALISIS_ARQUITECTURA.docx"
UML_PATH = Path(r"C:\Users\Usuario\Desktop\Trabajo Martin\UML\UML.png")


def set_run_font(run, name="Times New Roman", size=12, bold=False, italic=False, color=None):
    run.font.name = name
    run._element.rPr.rFonts.set(qn("w:eastAsia"), name)
    run.font.size = Pt(size)
    run.bold = bold
    run.italic = italic
    if color is not None:
        run.font.color.rgb = color


def set_paragraph_format(p, *, align=WD_ALIGN_PARAGRAPH.JUSTIFY, space_after=8, space_before=0, first_line=None, line=1.15):
    pf = p.paragraph_format
    p.alignment = align
    pf.space_after = Pt(space_after)
    pf.space_before = Pt(space_before)
    pf.line_spacing = line
    if first_line is not None:
        pf.first_line_indent = Cm(first_line)


def add_page_number(paragraph):
    run = paragraph.add_run()
    fldChar1 = OxmlElement("w:fldChar")
    fldChar1.set(qn("w:fldCharType"), "begin")
    instrText = OxmlElement("w:instrText")
    instrText.set(qn("xml:space"), "preserve")
    instrText.text = " PAGE "
    fldChar2 = OxmlElement("w:fldChar")
    fldChar2.set(qn("w:fldCharType"), "end")
    run._r.append(fldChar1)
    run._r.append(instrText)
    run._r.append(fldChar2)
    set_run_font(run, size=10)


def add_toc_field(paragraph):
    run = paragraph.add_run()
    r = run._r
    fld1 = OxmlElement("w:fldChar")
    fld1.set(qn("w:fldCharType"), "begin")
    instr = OxmlElement("w:instrText")
    instr.set(qn("xml:space"), "preserve")
    instr.text = ' TOC \\o "1-2" \\h \\z \\u '
    fld2 = OxmlElement("w:fldChar")
    fld2.set(qn("w:fldCharType"), "separate")
    text = OxmlElement("w:t")
    text.text = "Haga clic derecho sobre este campo y elija «Actualizar campo» para generar el índice."
    fld3 = OxmlElement("w:fldChar")
    fld3.set(qn("w:fldCharType"), "end")
    r.append(fld1)
    r.append(instr)
    r.append(fld2)
    r.append(text)
    r.append(fld3)
    set_run_font(run, size=11, italic=True)


def shade_cell(cell, color_hex):
    tc = cell._tePr if hasattr(cell, "_tePr") else cell._tc
    tcPr = tc.get_or_add_tcPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:fill"), color_hex)
    shd.set(qn("w:val"), "clear")
    tcPr.append(shd)


def set_cell_border(cell):
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()
    tcBorders = OxmlElement("w:tcBorders")
    for edge in ("top", "left", "bottom", "right"):
        el = OxmlElement(f"w:{edge}")
        el.set(qn("w:val"), "single")
        el.set(qn("w:sz"), "4")
        el.set(qn("w:space"), "0")
        el.set(qn("w:color"), "666666")
        tcBorders.append(el)
    tcPr.append(tcBorders)


def prevent_table_split(table):
    tbl = table._tbl
    tblPr = tbl.tblPr if tbl.tblPr is not None else OxmlElement("w:tblPr")
    tblW = OxmlElement("w:tblW")
    tblW.set(qn("w:w"), "5000")
    tblW.set(qn("w:type"), "pct")


def count_words(text: str) -> int:
    body = re.sub(r"```[\s\S]*?```", " ", text)
    body = re.sub(r"https?://\S+", " ", body)
    words = re.findall(r"[A-Za-zÁÉÍÓÚáéíóúÑñÜü0-9]+", body)
    return len(words)


def add_hyperlink(paragraph, text, url):
    part = paragraph.part
    r_id = part.relate_to(url, RT.HYPERLINK, is_external=True)
    hyperlink = OxmlElement("w:hyperlink")
    hyperlink.set(qn("r:id"), r_id)
    new_run = OxmlElement("w:r")
    rPr = OxmlElement("w:rPr")
    color = OxmlElement("w:color")
    color.set(qn("w:val"), "0000FF")
    u = OxmlElement("w:u")
    u.set(qn("w:val"), "single")
    rFonts = OxmlElement("w:rFonts")
    rFonts.set(qn("w:ascii"), "Times New Roman")
    rFonts.set(qn("w:hAnsi"), "Times New Roman")
    sz = OxmlElement("w:sz")
    sz.set(qn("w:val"), "24")
    rPr.append(rFonts)
    rPr.append(color)
    rPr.append(u)
    rPr.append(sz)
    new_run.append(rPr)
    t = OxmlElement("w:t")
    t.set(qn("xml:space"), "preserve")
    t.text = text
    new_run.append(t)
    hyperlink.append(new_run)
    paragraph._p.append(hyperlink)


def configure_styles(doc: Document):
    styles = doc.styles
    normal = styles["Normal"]
    normal.font.name = "Times New Roman"
    normal.font.size = Pt(12)
    normal._element.rPr.rFonts.set(qn("w:eastAsia"), "Times New Roman")
    pf = normal.paragraph_format
    pf.line_spacing = 1.15
    pf.space_after = Pt(8)
    pf.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY

    for i, size, space_before, space_after in (
        (1, 16, 18, 10),
        (2, 13, 14, 8),
        (3, 12, 12, 6),
    ):
        st = styles[f"Heading {i}"]
        st.font.name = "Times New Roman"
        st.font.size = Pt(size)
        st.font.bold = True
        st.font.color.rgb = RGBColor(0x1A, 0x1A, 0x1A)
        st._element.rPr.rFonts.set(qn("w:eastAsia"), "Times New Roman")
        st.paragraph_format.space_before = Pt(space_before)
        st.paragraph_format.space_after = Pt(space_after)
        st.paragraph_format.line_spacing = 1.15
        st.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.LEFT

    if "CodeBlock" not in [s.name for s in styles]:
        cs = styles.add_style("CodeBlock", WD_STYLE_TYPE.PARAGRAPH)
        cs.font.name = "Consolas"
        cs.font.size = Pt(9)
        cs._element.rPr.rFonts.set(qn("w:eastAsia"), "Consolas")
        cs.paragraph_format.line_spacing = 1.0
        cs.paragraph_format.space_after = Pt(2)
        cs.paragraph_format.space_before = Pt(0)
        cs.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.LEFT


def add_body_paragraph(doc, text, *, first_line=0.75):
    p = doc.add_paragraph()
    set_paragraph_format(p, first_line=first_line, space_after=8)
    add_inline_runs(p, text)
    return p


def add_inline_runs(paragraph, text, *, base_size=12, mono=False):
    """Parse **bold**, *italic* and `code` in a line."""
    pattern = re.compile(r"(\*\*[^*]+\*\*|`[^`]+`|\*[^*]+\*)")
    parts = pattern.split(text)
    for part in parts:
        if not part:
            continue
        if part.startswith("**") and part.endswith("**"):
            run = paragraph.add_run(part[2:-2])
            set_run_font(run, size=base_size, bold=True)
        elif part.startswith("`") and part.endswith("`"):
            run = paragraph.add_run(part[1:-1])
            set_run_font(run, name="Consolas", size=max(10, base_size - 1))
        elif part.startswith("*") and part.endswith("*"):
            run = paragraph.add_run(part[1:-1])
            set_run_font(run, size=base_size, italic=True)
        else:
            run = paragraph.add_run(part)
            if mono:
                set_run_font(run, name="Consolas", size=base_size)
            else:
                set_run_font(run, size=base_size)


def shade_paragraph(paragraph, fill="F4F4F4"):
    pPr = paragraph._p.get_or_add_pPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:val"), "clear")
    shd.set(qn("w:color"), "auto")
    shd.set(qn("w:fill"), fill)
    pPr.append(shd)


def add_code_block(doc, lines):
    for i, line in enumerate(lines):
        p = doc.add_paragraph()
        p.style = doc.styles["CodeBlock"]
        set_paragraph_format(p, align=WD_ALIGN_PARAGRAPH.LEFT, space_after=0, space_before=0, line=1.0)
        if i == 0:
            p.paragraph_format.space_before = Pt(8)
        if i == len(lines) - 1:
            p.paragraph_format.space_after = Pt(10)
        shade_paragraph(p, "F3F3F3")
        run = p.add_run(line if line else " ")
        set_run_font(run, name="Consolas", size=9)


def add_table(doc, rows):
    table = doc.add_table(rows=len(rows), cols=len(rows[0]))
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.autofit = True
    for i, row in enumerate(rows):
        for j, val in enumerate(row):
            cell = table.rows[i].cells[j]
            cell.text = ""
            p = cell.paragraphs[0]
            set_paragraph_format(p, align=WD_ALIGN_PARAGRAPH.LEFT, space_after=2, space_before=2, line=1.1)
            run = p.add_run(val)
            set_run_font(run, size=10, bold=(i == 0))
            set_cell_border(cell)
            if i == 0:
                shade_cell(cell, "E8E8E8")
    doc.add_paragraph()
    return table


def setup_sections(doc: Document):
    section = doc.sections[0]
    section.page_width = Cm(21.0)
    section.page_height = Cm(29.7)
    section.top_margin = Cm(2.5)
    section.bottom_margin = Cm(2.5)
    section.left_margin = Cm(2.5)
    section.right_margin = Cm(2.5)
    section.different_first_page_header_footer = True

    # First page: empty header/footer (cover)
    section.first_page_header.paragraphs[0].text = ""
    section.first_page_footer.paragraphs[0].text = ""

    header = section.header
    hp = header.paragraphs[0]
    hp.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    run = hp.add_run("Memoria del desarrollo del sistema Concesionario")
    set_run_font(run, size=9, italic=True, color=RGBColor(0x66, 0x66, 0x66))

    footer = section.footer
    fp = footer.paragraphs[0]
    fp.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = fp.add_run("Página ")
    set_run_font(run, size=10)
    add_page_number(fp)


def build_cover(doc: Document):
    for _ in range(3):
        p = doc.add_paragraph()
        set_paragraph_format(p, align=WD_ALIGN_PARAGRAPH.CENTER, space_after=0)

    p = doc.add_paragraph()
    set_paragraph_format(p, align=WD_ALIGN_PARAGRAPH.CENTER, space_after=6)
    run = p.add_run("UNIVERSIDAD INTERNACIONAL DE LA RIOJA (UNIR)")
    set_run_font(run, size=13, bold=True)

    p = doc.add_paragraph()
    set_paragraph_format(p, align=WD_ALIGN_PARAGRAPH.CENTER, space_after=28)
    run = p.add_run("Ingeniería Informática")
    set_run_font(run, size=12, italic=True)

    p = doc.add_paragraph()
    set_paragraph_format(p, align=WD_ALIGN_PARAGRAPH.CENTER, space_after=18)
    run = p.add_run("Memoria del desarrollo del sistema Concesionario")
    set_run_font(run, size=22, bold=True)

    p = doc.add_paragraph()
    set_paragraph_format(p, align=WD_ALIGN_PARAGRAPH.CENTER, space_after=36)
    run = p.add_run("Laboratorio de Programación Orientada a Objetos en Java")
    set_run_font(run, size=12, italic=True)

    fields = [
        ("Estudiante", "[NOMBRE DEL ESTUDIANTE]"),
        ("Asignatura", "[NOMBRE DE LA ASIGNATURA]"),
        ("Laboratorio", "[NÚMERO O NOMBRE DEL LABORATORIO]"),
        ("Fecha", "[FECHA]"),
    ]
    for label, value in fields:
        p = doc.add_paragraph()
        set_paragraph_format(p, align=WD_ALIGN_PARAGRAPH.CENTER, space_after=4)
        run = p.add_run(f"{label}: ")
        set_run_font(run, size=12, bold=True)
        run = p.add_run(value)
        set_run_font(run, size=12)

    p = doc.add_paragraph()
    set_paragraph_format(p, align=WD_ALIGN_PARAGRAPH.CENTER, space_before=40, space_after=0)
    run = p.add_run(
        "Valores sugeridos por el enunciado: asignatura Programación Avanzada; "
        "laboratorio #1 — Diseño e implementación de clases."
    )
    set_run_font(run, size=10, italic=True, color=RGBColor(0x55, 0x55, 0x55))

    # Nueva sección: índice y cuerpo, con numeración desde 1
    new_section = doc.add_section(WD_SECTION.NEW_PAGE)
    new_section.page_width = Cm(21.0)
    new_section.page_height = Cm(29.7)
    new_section.top_margin = Cm(2.5)
    new_section.bottom_margin = Cm(2.5)
    new_section.left_margin = Cm(2.5)
    new_section.right_margin = Cm(2.5)
    new_section.different_first_page_header_footer = False

    sectPr = new_section._sectPr
    pgNumType = OxmlElement("w:pgNumType")
    pgNumType.set(qn("w:start"), "1")
    sectPr.append(pgNumType)

    header = new_section.header
    header.is_linked_to_previous = False
    hp = header.paragraphs[0]
    hp.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    run = hp.add_run("Memoria del desarrollo del sistema Concesionario")
    set_run_font(run, size=9, italic=True, color=RGBColor(0x66, 0x66, 0x66))

    footer = new_section.footer
    footer.is_linked_to_previous = False
    fp = footer.paragraphs[0]
    fp.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = fp.add_run("Página ")
    set_run_font(run, size=10)
    add_page_number(fp)


def build_toc_page(doc: Document):
    p = doc.add_heading("Índice", level=1)
    p.alignment = WD_ALIGN_PARAGRAPH.LEFT

    note = doc.add_paragraph()
    set_paragraph_format(note, align=WD_ALIGN_PARAGRAPH.LEFT, space_after=10)
    run = note.add_run(
        "El índice automático se actualiza al abrir el documento en Microsoft Word "
        "(clic derecho sobre el campo → Actualizar campo → Actualizar toda la tabla). "
        "A continuación se incluye también un índice fijo."
    )
    set_run_font(run, size=10, italic=True)

    toc_p = doc.add_paragraph()
    set_paragraph_format(toc_p, align=WD_ALIGN_PARAGRAPH.LEFT, space_after=12)
    add_toc_field(toc_p)

    entries = [
        ("1. Introducción", "2"),
        ("1.1. Contexto académico", "2"),
        ("1.2. Problema que se pretende resolver", "2"),
        ("1.3. Objetivos generales y específicos", "2"),
        ("2. Análisis de requisitos", "3"),
        ("2.1. Requisitos extraídos del enunciado", "3"),
        ("2.2. Restricciones del proyecto", "3"),
        ("2.3. Funcionalidades principales", "3"),
        ("3. Proceso de desarrollo", "4"),
        ("4. Diseño orientado a objetos", "5"),
        ("5. Organización de la aplicación", "7"),
        ("6. Casos de uso", "7"),
        ("7. Decisiones y compensaciones del diseño", "8"),
        ("8. Comparación con BibliotecaPOO", "11"),
        ("9. Comprobación del UML", "11"),
        ("10. Limitaciones y recomendaciones", "12"),
        ("11. Pruebas propuestas", "12"),
        ("12. Conclusiones", "13"),
        ("13. Referencias", "13"),
    ]
    for title, page in entries:
        p = doc.add_paragraph()
        set_paragraph_format(p, align=WD_ALIGN_PARAGRAPH.LEFT, space_after=2, line=1.15)
        # dotted leaders via tab
        pPr = p._p.get_or_add_pPr()
        tabs = OxmlElement("w:tabs")
        tab = OxmlElement("w:tab")
        tab.set(qn("w:val"), "right")
        tab.set(qn("w:leader"), "dot")
        tab.set(qn("w:pos"), "9000")
        tabs.append(tab)
        pPr.append(tabs)
        run = p.add_run(title)
        set_run_font(run, size=12)
        run = p.add_run("\t" + page)
        set_run_font(run, size=12)

    doc.add_page_break()


def is_table_separator(line: str) -> bool:
    s = line.strip()
    return s.startswith("|") and set(s.replace("|", "").replace(":", "").replace("-", "").replace(" ", "")) == set()


def parse_table_row(line: str) -> list[str]:
    parts = [c.strip() for c in line.strip().strip("|").split("|")]
    return parts


def convert_md(doc: Document, md: str):
    lines = md.splitlines()
    i = 0
    # skip cover/index already in MD (we rebuild them)
    # Find start at "## 1. Introducción"
    start = 0
    for idx, line in enumerate(lines):
        if line.startswith("## 1. Introducción"):
            start = idx
            break
    i = start
    in_code = False
    code_lines: list[str] = []

    while i < len(lines):
        line = lines[i]

        if line.startswith("```"):
            if not in_code:
                in_code = True
                code_lines = []
            else:
                add_code_block(doc, code_lines)
                in_code = False
                code_lines = []
            i += 1
            continue

        if in_code:
            code_lines.append(line)
            i += 1
            continue

        if not line.strip():
            i += 1
            continue

        if line.startswith("## "):
            text = line[3:].strip()
            doc.add_heading(text, level=1)
            i += 1
            continue

        if line.startswith("### "):
            text = line[4:].strip()
            doc.add_heading(text, level=2)
            i += 1
            continue

        if line.startswith("**Figura 1.**"):
            if UML_PATH.exists():
                p = doc.add_paragraph()
                set_paragraph_format(p, align=WD_ALIGN_PARAGRAPH.CENTER, space_before=8, space_after=4)
                run = p.add_run()
                run.add_picture(str(UML_PATH), width=Cm(16.0))
            cap = doc.add_paragraph()
            set_paragraph_format(cap, align=WD_ALIGN_PARAGRAPH.CENTER, space_after=12)
            run = cap.add_run(
                "Figura 1. Diagrama de clases UML del sistema Concesionario (archivo UML/UML.png)."
            )
            set_run_font(run, size=10, italic=True)
            i += 1
            continue

        if line.strip().startswith("|") and i + 1 < len(lines) and is_table_separator(lines[i + 1]):
            rows = [parse_table_row(line)]
            i += 2
            while i < len(lines) and lines[i].strip().startswith("|"):
                rows.append(parse_table_row(lines[i]))
                i += 1
            add_table(doc, rows)
            continue

        if re.match(r"^- ", line):
            items = []
            while i < len(lines) and re.match(r"^- ", lines[i]):
                items.append(lines[i][2:].strip())
                i += 1
            for item in items:
                p = doc.add_paragraph(style="List Bullet")
                set_paragraph_format(p, align=WD_ALIGN_PARAGRAPH.JUSTIFY, space_after=4, first_line=0)
                add_inline_runs(p, item)
            continue

        if re.match(r"^\d+\. ", line):
            items = []
            while i < len(lines) and re.match(r"^\d+\. ", lines[i]):
                items.append(re.sub(r"^\d+\. ", "", lines[i]).strip())
                i += 1
            for item in items:
                p = doc.add_paragraph(style="List Number")
                set_paragraph_format(p, align=WD_ALIGN_PARAGRAPH.JUSTIFY, space_after=4, first_line=0)
                add_inline_runs(p, item)
            continue

        # normal paragraph: join wrapped? MD is already one paragraph per line block
        # Consecutive non-empty non-special lines are separate paragraphs in our MD
        add_body_paragraph(doc, line.strip())
        i += 1


def main():
    md = MD_PATH.read_text(encoding="utf-8")
    body = md.split("## 1. Introducción", 1)[-1]
    body = "## 1. Introducción" + body
    words = count_words(body)
    print(f"Palabras del cuerpo (secciones 1-13): {words}")

    doc = Document()
    configure_styles(doc)
    setup_sections(doc)
    build_cover(doc)
    build_toc_page(doc)
    convert_md(doc, md)

    pages_content = round(words / 420, 1)
    pages_total = round(pages_content + 2, 1)  # portada + índice
    print(f"Páginas estimadas de contenido: {pages_content}")
    print(f"Páginas totales estimadas (con portada e índice): {pages_total}")

    doc.save(DOCX_PATH)
    print(f"Escrito: {DOCX_PATH}")


if __name__ == "__main__":
    main()
